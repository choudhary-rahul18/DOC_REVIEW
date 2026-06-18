import logging
import re
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)


def parse_document(file_path: str | Path) -> str:
    """
    Dispatch to the correct parser by file extension.
    Returns raw text with whitespace lightly normalised.
    Raises ValueError for unsupported extensions.
    Raises RuntimeError (wrapping the library error) on parse failure.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {list(_PARSERS.keys())}"
        )
    try:
        raw = parser(file_path)
    except Exception as e:
        raise RuntimeError(f"Failed to parse {file_path.name}: {e}") from e

    logger.info("Parsed %s: %d chars", file_path.name, len(raw))
    return _normalise(raw)


# ── PDF ────────────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: Path) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""  # None on image-only pages
            if not text.strip():
                logger.warning(
                    "Page %d of %s yielded no text (image-only or empty)",
                    i + 1, file_path.name,
                )
            pages.append(text)

    return "\n".join(pages)


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: Path) -> str:
    from docx import Document

    doc = Document(file_path)
    # Join with a space, not newline.
    # DOCX files often split a single speaker-turn across multiple doc.paragraphs
    # entries due to soft returns or style breaks. Joining with "\n" would cause
    # the chunker's speaker-discovery regex to see each fragment in isolation,
    # missing speaker labels split at paragraph boundaries.
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    return " ".join(parts)


# ── Normalisation ──────────────────────────────────────────────────────────────

def _normalise(text: str) -> str:
    # 1. Unicode: curly quotes/apostrophes → ASCII; em/en dashes → hyphen;
    #    non-breaking space → regular space.
    text = text.replace('‘', "'").replace('’', "'")
    text = text.replace('“', '"').replace('”', '"')
    text = text.replace('–', '-').replace('—', '-')
    text = text.replace(' ', ' ')

    # 2. Reconstruct soft-hyphenated words split across lines ("immuniza-\ntion" → "immunization").
    text = re.sub(r'-\n(\S)', lambda m: m.group(1), text)

    # 3. Collapse soft line breaks — a newline NOT preceded by sentence-ending
    #    punctuation and NOT followed by another newline is a PDF word-wrap
    #    artifact, not a paragraph boundary. Replace with a space so the
    #    sentence stays intact for the chunker's sentence splitter.
    #    Sentence-ending punctuation: . ! ? … and their quoted variants ." !'
    text = re.sub(r'(?<![.!?…"\'»])\n(?!\n)', ' ', text)

    # 4. Collapse 2+ non-newline whitespace (tabs, multiple spaces) → single space.
    text = re.sub(r'[^\S\n]{2,}', ' ', text)

    # 5. Collapse 3+ consecutive newlines → double newline (paragraph separator).
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


# ── Dispatch ───────────────────────────────────────────────────────────────────

_PARSERS: dict[str, Callable[[Path], str]] = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
}
